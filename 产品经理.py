import sys
import pandas as pd
import numpy as np
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sqlite3
import json
import requests
from bs4 import BeautifulSoup
import threading
import time
from sklearn.linear_model import LinearRegression
from sklearn.ensemble import RandomForestRegressor
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
import warnings
warnings.filterwarnings('ignore')

from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QPushButton, QTabWidget, QTableWidget, QTableWidgetItem, 
                             QTreeWidget, QTreeWidgetItem, QSplitter, QTextEdit, QComboBox,
                             QLineEdit, QLabel, QDateEdit, QSpinBox, QDoubleSpinBox,
                             QMessageBox, QFileDialog, QProgressBar, QStatusBar, QToolBar,
                             QAction, QDockWidget, QListWidget, QHeaderView, QFormLayout,
                             QGroupBox, QCheckBox, QRadioButton, QButtonGroup, QSlider,
                             QDialog, QDialogButtonBox, QGridLayout, QStackedWidget, 
                             QFrame, QSizePolicy, QMenu, QSystemTrayIcon, QStyle, QInputDialog)
from PyQt5.QtCore import Qt, QDate, QSize, QThread, pyqtSignal, QTimer, QSettings, QPoint, QModelIndex
from PyQt5.QtGui import QIcon, QFont, QPalette, QColor, QPainter, QLinearGradient, QBrush, QPen, QPixmap, QMovie


# 数据库管理类 - 增强版
class DatabaseManager:
    def __init__(self, db_name="product_manager.db"):
        self.db_name = db_name
        self.init_db()
    
    def init_db(self):
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        
        # 创建项目表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS projects (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                description TEXT,
                status TEXT,
                priority INTEGER,
                start_date TEXT,
                end_date TEXT,
                budget REAL,
                actual_cost REAL,
                progress INTEGER,
                created_date TEXT,
                last_modified TEXT,
                tags TEXT,
                risk_level TEXT,
                stakeholders TEXT
            )
        ''')
        
        # 创建团队成员表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS team_members (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                role TEXT,
                department TEXT,
                skills TEXT,
                workload INTEGER,
                performance_score REAL,
                contact_info TEXT,
                hourly_rate REAL,
                available_hours INTEGER
            )
        ''')
        
        # 创建任务表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS tasks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                project_id INTEGER,
                name TEXT NOT NULL,
                description TEXT,
                assignee_id INTEGER,
                status TEXT,
                priority INTEGER,
                estimated_hours REAL,
                actual_hours REAL,
                start_date TEXT,
                due_date TEXT,
                dependencies TEXT,
                tags TEXT,
                FOREIGN KEY (project_id) REFERENCES projects (id),
                FOREIGN KEY (assignee_id) REFERENCES team_members (id)
            )
        ''')
        
        # 创建指标表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS metrics (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                project_id INTEGER,
                date TEXT,
                kpi_name TEXT,
                value REAL,
                target REAL,
                notes TEXT,
                FOREIGN KEY (project_id) REFERENCES projects (id)
            )
        ''')
        
        # 创建资源表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS resources (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                type TEXT,
                cost REAL,
                availability TEXT,
                assigned_to INTEGER,
                project_id INTEGER,
                FOREIGN KEY (assigned_to) REFERENCES team_members (id),
                FOREIGN KEY (project_id) REFERENCES projects (id)
            )
        ''')
        
        # 创建风险评估表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS risks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                project_id INTEGER,
                description TEXT,
                impact TEXT,
                probability TEXT,
                mitigation_plan TEXT,
                status TEXT,
                assigned_to INTEGER,
                due_date TEXT,
                FOREIGN KEY (project_id) REFERENCES projects (id),
                FOREIGN KEY (assigned_to) REFERENCES team_members (id)
            )
        ''')
        
        # 创建文档表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                project_id INTEGER,
                name TEXT NOT NULL,
                type TEXT,
                path TEXT,
                upload_date TEXT,
                version TEXT,
                uploaded_by INTEGER,
                FOREIGN KEY (project_id) REFERENCES projects (id),
                FOREIGN KEY (uploaded_by) REFERENCES team_members (id)
            )
        ''')
        
        conn.commit()
        conn.close()
    
    def execute_query(self, query, params=None, fetch=True):
        conn = sqlite3.connect(self.db_name)
        cursor = conn.cursor()
        try:
            if params:
                cursor.execute(query, params)
            else:
                cursor.execute(query)
            conn.commit()
            if fetch:
                result = cursor.fetchall()
            else:
                result = None
        except Exception as e:
            result = f"Error: {str(e)}"
        finally:
            conn.close()
        return result
    
    def get_projects(self, status=None):
        if status:
            return self.execute_query("SELECT * FROM projects WHERE status=?", (status,))
        return self.execute_query("SELECT * FROM projects")
    
    def get_project(self, project_id):
        return self.execute_query("SELECT * FROM projects WHERE id=?", (project_id,))
    
    def get_team_members(self, department=None):
        if department:
            return self.execute_query("SELECT * FROM team_members WHERE department=?", (department,))
        return self.execute_query("SELECT * FROM team_members")
    
    def get_tasks(self, project_id=None, assignee_id=None):
        if project_id and assignee_id:
            return self.execute_query("SELECT * FROM tasks WHERE project_id=? AND assignee_id=?", (project_id, assignee_id))
        elif project_id:
            return self.execute_query("SELECT * FROM tasks WHERE project_id=?", (project_id,))
        elif assignee_id:
            return self.execute_query("SELECT * FROM tasks WHERE assignee_id=?", (assignee_id,))
        return self.execute_query("SELECT * FROM tasks")
    
    def get_metrics(self, project_id=None, kpi_name=None):
        if project_id and kpi_name:
            return self.execute_query("SELECT * FROM metrics WHERE project_id=? AND kpi_name=?", (project_id, kpi_name))
        elif project_id:
            return self.execute_query("SELECT * FROM metrics WHERE project_id=?", (project_id,))
        elif kpi_name:
            return self.execute_query("SELECT * FROM metrics WHERE kpi_name=?", (kpi_name,))
        return self.execute_query("SELECT * FROM metrics")
    
    def get_risks(self, project_id=None):
        if project_id:
            return self.execute_query("SELECT * FROM risks WHERE project_id=?", (project_id,))
        return self.execute_query("SELECT * FROM risks")
    
    def get_resources(self, project_id=None):
        if project_id:
            return self.execute_query("SELECT * FROM resources WHERE project_id=?", (project_id,))
        return self.execute_query("SELECT * FROM resources")
    
    def get_documents(self, project_id=None):
        if project_id:
            return self.execute_query("SELECT * FROM documents WHERE project_id=?", (project_id,))
        return self.execute_query("SELECT * FROM documents")
    
    def backup_database(self, backup_path):
        try:
            import shutil
            shutil.copy2(self.db_name, backup_path)
            return True
        except Exception as e:
            return False
    
    def restore_database(self, backup_path):
        try:
            import shutil
            shutil.copy2(backup_path, self.db_name)
            return True
        except Exception as e:
            return False


# 高级图表组件
class AdvancedMplCanvas(FigureCanvas):
    def __init__(self, parent=None, width=5, height=4, dpi=100):
        self.fig = Figure(figsize=(width, height), dpi=dpi)
        super().__init__(self.fig)
        self.setParent(parent)
        
    def plot_line(self, data, title="", xlabel="", ylabel="", legend=None):
        self.fig.clear()
        ax = self.fig.add_subplot(111)
        
        if isinstance(data, dict):
            for key, values in data.items():
                ax.plot(values['x'], values['y'], label=key)
        else:
            ax.plot(data['x'], data['y'])
        
        ax.set_title(title)
        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        
        if legend:
            ax.legend()
        
        self.fig.tight_layout()
        self.draw()
    
    def plot_bar(self, data, title="", xlabel="", ylabel="", stacked=False):
        self.fig.clear()
        ax = self.fig.add_subplot(111)
        
        # 处理单组数据的情况
        if not isinstance(data, dict):
            # 如果是单组数据，转换为字典格式
            data = {'Data': data}
        
        if stacked:
            bottom = None
            for i, (key, values) in enumerate(data.items()):
                if isinstance(values, dict) and 'x' in values and 'y' in values:
                    x_data = values['x']
                    y_data = values['y']
                else:
                    # 如果不是字典格式，直接使用值
                    x_data = range(len(values))
                    y_data = values
                
                if bottom is None:
                    ax.bar(x_data, y_data, label=key)
                    bottom = y_data
                else:
                    ax.bar(x_data, y_data, bottom=bottom, label=key)
                    bottom = [b + v for b, v in zip(bottom, y_data)]
        else:
            # 获取所有组的长度
            max_length = 0
            for values in data.values():
                if isinstance(values, dict) and 'x' in values and 'y' in values:
                    max_length = max(max_length, len(values['y']))
                else:
                    max_length = max(max_length, len(values))
            
            width = 0.8 / len(data)
            for i, (key, values) in enumerate(data.items()):
                if isinstance(values, dict) and 'x' in values and 'y' in values:
                    x_data = values['x']
                    y_data = values['y']
                else:
                    x_data = range(len(values))
                    y_data = values
                
                # 确保所有组有相同的长度
                if len(y_data) < max_length:
                    y_data = list(y_data) + [0] * (max_length - len(y_data))
                
                x_pos = [p + i * width for p in range(len(y_data))]
                ax.bar(x_pos, y_data, width=width, label=key)
            
            ax.set_xticks([p + (len(data) - 1) * width / 2 for p in range(max_length)])
            if isinstance(next(iter(data.values())), dict) and 'x' in next(iter(data.values())):
                ax.set_xticklabels(next(iter(data.values()))['x'])
            else:
                ax.set_xticklabels(range(max_length))
        
        ax.set_title(title)
        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        ax.legend()
        
        self.fig.tight_layout()
        self.draw()
    
    def plot_pie(self, data, title=""):
        self.fig.clear()
        ax = self.fig.add_subplot(111)
        
        ax.pie(data['y'], labels=data['x'], autopct='%1.1f%%')
        ax.set_title(title)
        
        self.fig.tight_layout()
        self.draw()
    
    def plot_heatmap(self, data, title="", xlabel="", ylabel=""):
        self.fig.clear()
        ax = self.fig.add_subplot(111)
        
        sns.heatmap(data, annot=True, fmt=".2f", ax=ax)
        ax.set_title(title)
        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        
        self.fig.tight_layout()
        self.draw()
    
    def plot_scatter(self, data, title="", xlabel="", ylabel=""):
        self.fig.clear()
        ax = self.fig.add_subplot(111)
        
        ax.scatter(data['x'], data['y'])
        ax.set_title(title)
        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        
        # Add trend line
        z = np.polyfit(data['x'], data['y'], 1)
        p = np.poly1d(z)
        ax.plot(data['x'], p(data['x']), "r--")
        
        self.fig.tight_layout()
        self.draw()
    
    def plot_histogram(self, data, title="", xlabel="", ylabel="", bins=10):
        self.fig.clear()
        ax = self.fig.add_subplot(111)
        
        ax.hist(data, bins=bins)
        ax.set_title(title)
        ax.set_xlabel(xlabel)
        ax.set_ylabel(ylabel)
        
        self.fig.tight_layout()
        self.draw()
    
    def clear_plot(self):
        self.fig.clear()
        self.draw()


# 甘特图组件
class GanttWidget(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()
        self.tasks = []
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # 控制按钮
        control_layout = QHBoxLayout()
        self.zoom_in_btn = QPushButton("Zoom In")
        self.zoom_out_btn = QPushButton("Zoom Out")
        self.today_btn = QPushButton("Today")
        control_layout.addWidget(self.zoom_in_btn)
        control_layout.addWidget(self.zoom_out_btn)
        control_layout.addWidget(self.today_btn)
        control_layout.addStretch()
        
        layout.addLayout(control_layout)
        
        # 甘特图画布
        self.canvas = AdvancedMplCanvas(self, width=10, height=6)
        layout.addWidget(self.canvas)
        
        # 连接信号
        self.zoom_in_btn.clicked.connect(self.zoom_in)
        self.zoom_out_btn.clicked.connect(self.zoom_out)
        self.today_btn.clicked.connect(self.show_today)
    
    def set_tasks(self, tasks):
        self.tasks = tasks
        self.update_gantt()
    
    def update_gantt(self):
        if not self.tasks:
            return
        
        # 准备数据
        task_names = []
        start_dates = []
        end_dates = []
        colors = []
        
        for task in self.tasks:
            task_names.append(task['name'])
            start_dates.append(datetime.strptime(task['start_date'], '%Y-%m-%d'))
            end_dates.append(datetime.strptime(task['end_date'], '%Y-%m-%d'))
            
            # 根据状态设置颜色
            if task['status'] == 'Completed':
                colors.append('green')
            elif task['status'] == 'In Progress':
                colors.append('blue')
            elif task['status'] == 'Delayed':
                colors.append('red')
            else:
                colors.append('gray')
        
        # 绘制甘特图
        self.canvas.fig.clear()
        ax = self.canvas.fig.add_subplot(111)
        
        for i, (task, start, end, color) in enumerate(zip(task_names, start_dates, end_dates, colors)):
            duration = (end - start).days
            ax.barh(task, duration, left=start, color=color, alpha=0.6)
        
        ax.set_xlabel('Timeline')
        ax.set_ylabel('Tasks')
        ax.set_title('Project Gantt Chart')
        
        # 格式化日期
        ax.xaxis_date()
        self.canvas.fig.autofmt_xdate()
        
        self.canvas.fig.tight_layout()
        self.canvas.draw()
    
    def zoom_in(self):
        # 实现缩放功能
        pass
    
    def zoom_out(self):
        # 实现缩放功能
        pass
    
    def show_today(self):
        # 显示今天的位置
        pass


# 高级数据分析线程
class AdvancedAnalysisThread(QThread):
    progress_signal = pyqtSignal(int, str)
    result_signal = pyqtSignal(object, str)
    
    def __init__(self, data, analysis_type, params=None):
        super().__init__()
        self.data = data
        self.analysis_type = analysis_type
        self.params = params or {}
    
    def run(self):
        try:
            if self.analysis_type == "forecast":
                self.progress_signal.emit(20, "Preparing data for forecasting...")
                # 使用机器学习进行预测
                X = np.array(range(len(self.data))).reshape(-1, 1)
                y = self.data['value'].values
                
                self.progress_signal.emit(40, "Training forecasting model...")
                model = RandomForestRegressor(n_estimators=100, random_state=42)
                model.fit(X, y)
                
                self.progress_signal.emit(70, "Generating forecast...")
                future_periods = self.params.get('periods', 30)
                future_X = np.array(range(len(self.data), len(self.data) + future_periods)).reshape(-1, 1)
                forecast = model.predict(future_X)
                
                result = {
                    'forecast': forecast,
                    'model_score': model.score(X, y),
                    'dates': pd.date_range(start=self.data['date'].iloc[-1], periods=future_periods+1, freq='D')[1:]
                }
                self.progress_signal.emit(100, "Forecast completed!")
                self.result_signal.emit(result, "forecast")
            
            elif self.analysis_type == "clustering":
                self.progress_signal.emit(20, "Preparing data for clustering...")
                # 使用K-means进行聚类分析
                features = self.data[['feature1', 'feature2']].values
                
                self.progress_signal.emit(40, "Scaling features...")
                scaler = StandardScaler()
                features_scaled = scaler.fit_transform(features)
                
                self.progress_signal.emit(60, "Performing clustering...")
                n_clusters = self.params.get('n_clusters', 3)
                kmeans = KMeans(n_clusters=n_clusters, random_state=42)
                clusters = kmeans.fit_predict(features_scaled)
                
                result = {
                    'clusters': clusters,
                    'centers': kmeans.cluster_centers_,
                    'inertia': kmeans.inertia_,
                    'labels': kmeans.labels_
                }
                self.progress_signal.emit(100, "Clustering completed!")
                self.result_signal.emit(result, "clustering")
            
            elif self.analysis_type == "sentiment":
                self.progress_signal.emit(20, "Analyzing sentiment...")
                # 简单的情感分析实现
                # 这里使用模拟实现，实际应用中可以使用NLTK或TextBlob等库
                sentiments = []
                for text in self.data['text']:
                    # 简单的情感分析逻辑
                    positive_words = ['good', 'great', 'excellent', 'amazing', 'wonderful', 'positive']
                    negative_words = ['bad', 'terrible', 'awful', 'horrible', 'negative', 'poor']
                    
                    positive_count = sum(1 for word in text.split() if word.lower() in positive_words)
                    negative_count = sum(1 for word in text.split() if word.lower() in negative_words)
                    
                    if positive_count > negative_count:
                        sentiments.append('Positive')
                    elif negative_count > positive_count:
                        sentiments.append('Negative')
                    else:
                        sentiments.append('Neutral')
                
                result = pd.DataFrame({
                    'text': self.data['text'],
                    'sentiment': sentiments
                })
                self.progress_signal.emit(100, "Sentiment analysis completed!")
                self.result_signal.emit(result, "sentiment")
        
        except Exception as e:
            self.result_signal.emit(f"Error: {str(e)}", "error")


# 实时数据监视器
class RealTimeMonitor(QThread):
    data_updated = pyqtSignal(dict)
    
    def __init__(self, update_interval=10):
        super().__init__()
        self.update_interval = update_interval
        self.running = False
    
    def run(self):
        self.running = True
        while self.running:
            # 模拟获取实时数据
            new_data = {
                'timestamp': datetime.now().isoformat(),
                'active_users': np.random.randint(100, 1000),
                'server_load': np.random.uniform(0.1, 0.9),
                'response_time': np.random.uniform(50, 500),
                'errors': np.random.randint(0, 10)
            }
            
            self.data_updated.emit(new_data)
            time.sleep(self.update_interval)
    
    def stop(self):
        self.running = False


# 项目详情面板 - 增强版
class AdvancedProjectDetailsWidget(QWidget):
    def __init__(self, project_data=None, db_manager=None):
        super().__init__()
        self.project_data = project_data or {}
        self.db_manager = db_manager
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        # 基本信息组
        basic_info_group = QGroupBox("Basic Information")
        basic_layout = QFormLayout()
        
        self.name_edit = QLineEdit()
        self.desc_edit = QTextEdit()
        self.status_combo = QComboBox()
        self.status_combo.addItems(["Planning", "Active", "On Hold", "Completed", "Cancelled"])
        self.priority_spin = QSpinBox()
        self.priority_spin.setRange(1, 5)
        self.risk_combo = QComboBox()
        self.risk_combo.addItems(["Low", "Medium", "High", "Critical"])
        
        basic_layout.addRow("Project Name:", self.name_edit)
        basic_layout.addRow("Description:", self.desc_edit)
        basic_layout.addRow("Status:", self.status_combo)
        basic_layout.addRow("Priority:", self.priority_spin)
        basic_layout.addRow("Risk Level:", self.risk_combo)
        
        basic_info_group.setLayout(basic_layout)
        layout.addWidget(basic_info_group)
        
        # 时间信息组
        time_info_group = QGroupBox("Timeline")
        time_layout = QFormLayout()
        
        self.start_date_edit = QDateEdit()
        self.start_date_edit.setCalendarPopup(True)
        self.end_date_edit = QDateEdit()
        self.end_date_edit.setCalendarPopup(True)
        self.progress_slider = QSlider(Qt.Horizontal)
        self.progress_slider.setRange(0, 100)
        self.progress_label = QLabel("0%")
        
        time_layout.addRow("Start Date:", self.start_date_edit)
        time_layout.addRow("End Date:", self.end_date_edit)
        time_layout.addRow("Progress:", self.progress_slider)
        time_layout.addRow("", self.progress_label)
        
        time_info_group.setLayout(time_layout)
        layout.addWidget(time_info_group)
        
        # 财务信息组
        financial_group = QGroupBox("Financial Information")
        financial_layout = QFormLayout()
        
        self.budget_spin = QDoubleSpinBox()
        self.budget_spin.setRange(0, 10000000)
        self.budget_spin.setPrefix("$ ")
        self.actual_cost_spin = QDoubleSpinBox()
        self.actual_cost_spin.setRange(0, 10000000)
        self.actual_cost_spin.setPrefix("$ ")
        
        financial_layout.addRow("Budget:", self.budget_spin)
        financial_layout.addRow("Actual Cost:", self.actual_cost_spin)
        
        financial_group.setLayout(financial_layout)
        layout.addWidget(financial_group)
        
        # 标签和利益相关者
        meta_group = QGroupBox("Additional Information")
        meta_layout = QFormLayout()
        
        self.tags_edit = QLineEdit()
        self.tags_edit.setPlaceholderText("Comma-separated tags")
        self.stakeholders_edit = QTextEdit()
        self.stakeholders_edit.setPlaceholderText("List of stakeholders, one per line")
        
        meta_layout.addRow("Tags:", self.tags_edit)
        meta_layout.addRow("Stakeholders:", self.stakeholders_edit)
        
        meta_group.setLayout(meta_layout)
        layout.addWidget(meta_group)
        
        self.setLayout(layout)
        
        # 连接信号
        self.progress_slider.valueChanged.connect(self.update_progress_label)
        
        if self.project_data:
            self.load_data()
    
    def update_progress_label(self, value):
        self.progress_label.setText(f"{value}%")
    
    def load_data(self):
        if self.project_data:
            self.name_edit.setText(self.project_data.get('name', ''))
            self.desc_edit.setPlainText(self.project_data.get('description', ''))
            self.status_combo.setCurrentText(self.project_data.get('status', 'Planning'))
            self.priority_spin.setValue(self.project_data.get('priority', 1))
            self.risk_combo.setCurrentText(self.project_data.get('risk_level', 'Low'))
            
            start_date = QDate.fromString(self.project_data.get('start_date', ''), Qt.ISODate)
            if start_date.isValid():
                self.start_date_edit.setDate(start_date)
            
            end_date = QDate.fromString(self.project_data.get('end_date', ''), Qt.ISODate)
            if end_date.isValid():
                self.end_date_edit.setDate(end_date)
            
            self.progress_slider.setValue(self.project_data.get('progress', 0))
            self.budget_spin.setValue(self.project_data.get('budget', 0))
            self.actual_cost_spin.setValue(self.project_data.get('actual_cost', 0))
            
            self.tags_edit.setText(self.project_data.get('tags', ''))
            self.stakeholders_edit.setPlainText(self.project_data.get('stakeholders', ''))
    
    def get_data(self):
        return {
            'name': self.name_edit.text(),
            'description': self.desc_edit.toPlainText(),
            'status': self.status_combo.currentText(),
            'priority': self.priority_spin.value(),
            'risk_level': self.risk_combo.currentText(),
            'start_date': self.start_date_edit.date().toString(Qt.ISODate),
            'end_date': self.end_date_edit.date().toString(Qt.ISODate),
            'progress': self.progress_slider.value(),
            'budget': self.budget_spin.value(),
            'actual_cost': self.actual_cost_spin.value(),
            'tags': self.tags_edit.text(),
            'stakeholders': self.stakeholders_edit.toPlainText()
        }


# 资源管理对话框
class ResourceManagerDialog(QDialog):
    def __init__(self, db_manager, parent=None):
        super().__init__(parent)
        self.db_manager = db_manager
        self.setWindowTitle("Resource Manager")
        self.setModal(True)
        self.setMinimumSize(800, 600)
        self.init_ui()
        self.load_data()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        # 工具栏
        toolbar = QToolBar()
        add_action = QAction("Add Resource", self)
        edit_action = QAction("Edit Resource", self)
        delete_action = QAction("Delete Resource", self)
        
        toolbar.addAction(add_action)
        toolbar.addAction(edit_action)
        toolbar.addAction(delete_action)
        
        layout.addWidget(toolbar)
        
        # 资源表格
        self.resource_table = QTableWidget()
        self.resource_table.setColumnCount(7)
        self.resource_table.setHorizontalHeaderLabels([
            "ID", "Name", "Type", "Cost", "Availability", "Assigned To", "Project"
        ])
        self.resource_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        
        layout.addWidget(self.resource_table)
        
        # 按钮框
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        
        layout.addWidget(button_box)
        
        self.setLayout(layout)
        
        # 连接信号
        add_action.triggered.connect(self.add_resource)
        edit_action.triggered.connect(self.edit_resource)
        delete_action.triggered.connect(self.delete_resource)
    
    def load_data(self):
        resources = self.db_manager.get_resources()
        self.resource_table.setRowCount(len(resources))
        
        for row, resource in enumerate(resources):
            for col, value in enumerate(resource):
                self.resource_table.setItem(row, col, QTableWidgetItem(str(value)))
    
    def add_resource(self):
        # 实现添加资源功能
        pass
    
    def edit_resource(self):
        # 实现编辑资源功能
        pass
    
    def delete_resource(self):
        # 实现删除资源功能
        pass


# 风险管理对话框
class RiskManagerDialog(QDialog):
    def __init__(self, db_manager, project_id=None, parent=None):
        super().__init__(parent)
        self.db_manager = db_manager
        self.project_id = project_id
        self.setWindowTitle("Risk Manager")
        self.setModal(True)
        self.setMinimumSize(800, 600)
        self.init_ui()
        self.load_data()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        # 工具栏
        toolbar = QToolBar()
        add_action = QAction("Add Risk", self)
        edit_action = QAction("Edit Risk", self)
        delete_action = QAction("Delete Risk", self)
        
        toolbar.addAction(add_action)
        toolbar.addAction(edit_action)
        toolbar.addAction(delete_action)
        
        layout.addWidget(toolbar)
        
        # 风险表格
        self.risk_table = QTableWidget()
        self.risk_table.setColumnCount(9)
        self.risk_table.setHorizontalHeaderLabels([
            "ID", "Project", "Description", "Impact", "Probability", "Mitigation Plan", "Status", "Assigned To", "Due Date"
        ])
        self.risk_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        
        layout.addWidget(self.risk_table)
        
        # 风险矩阵图
        risk_matrix_group = QGroupBox("Risk Matrix")
        risk_matrix_layout = QHBoxLayout()
        
        self.risk_matrix_canvas = AdvancedMplCanvas(self, width=5, height=4)
        risk_matrix_layout.addWidget(self.risk_matrix_canvas)
        
        risk_matrix_group.setLayout(risk_matrix_layout)
        layout.addWidget(risk_matrix_group)
        
        # 按钮框
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        
        layout.addWidget(button_box)
        
        self.setLayout(layout)
        
        # 连接信号
        add_action.triggered.connect(self.add_risk)
        edit_action.triggered.connect(self.edit_risk)
        delete_action.triggered.connect(self.delete_risk)
        
        # 绘制风险矩阵
        self.plot_risk_matrix()
    
    def load_data(self):
        if self.project_id:
            risks = self.db_manager.get_risks(self.project_id)
        else:
            risks = self.db_manager.get_risks()
        
        self.risk_table.setRowCount(len(risks))
        
        for row, risk in enumerate(risks):
            for col, value in enumerate(risk):
                self.risk_table.setItem(row, col, QTableWidgetItem(str(value)))
    
    def plot_risk_matrix(self):
        # 绘制风险矩阵图
        impact_levels = ['Low', 'Medium', 'High', 'Critical']
        probability_levels = ['Low', 'Medium', 'High']
        
        # 模拟风险数据
        risk_data = {
            'Low': [5, 3, 1],
            'Medium': [3, 2, 0],
            'High': [1, 0, 0],
            'Critical': [0, 0, 0]
        }
        
        self.risk_matrix_canvas.plot_heatmap(
            pd.DataFrame(risk_data, index=probability_levels),
            title="Risk Matrix",
            xlabel="Impact",
            ylabel="Probability"
        )
    
    def add_risk(self):
        # 实现添加风险功能
        pass
    
    def edit_risk(self):
        # 实现编辑风险功能
        pass
    
    def delete_risk(self):
        # 实现删除风险功能
        pass


# 高级产品经理总监系统主窗口
class AdvancedProductManagerDirectorSystem(QMainWindow):
    def __init__(self):
        super().__init__()
        self.db_manager = DatabaseManager()
        self.current_project_id = None
        self.settings = QSettings("ProductManagement", "DirectorSystem")
        self.init_ui()
        self.load_settings()
        
        # 启动实时数据监视器
        self.monitor = RealTimeMonitor(update_interval=5)
        self.monitor.data_updated.connect(self.update_real_time_data)
        self.monitor.start()
    
    def init_ui(self):
        self.setWindowTitle("Advanced Product Manager Director System")
        self.setGeometry(100, 100, 1400, 900)
        
        # 设置应用程序图标
        self.setWindowIcon(QIcon("icons/app_icon.png"))
        
        # 创建中央部件和选项卡
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout(central_widget)
        
        self.tab_widget = QTabWidget()
        self.tab_widget.setTabPosition(QTabWidget.West)
        self.tab_widget.setMovable(True)
        main_layout.addWidget(self.tab_widget)
        
        # 创建各个功能选项卡
        self.create_dashboard_tab()
        self.create_projects_tab()
        self.create_team_tab()
        self.create_analytics_tab()
        self.create_reports_tab()
        self.create_risks_tab()
        self.create_resources_tab()
        self.create_documents_tab()
        
        # 创建菜单栏
        self.create_menubar()
        
        # 创建工具栏
        self.create_toolbar()
        
        # 创建状态栏
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        
        # 添加状态栏组件
        self.status_label = QLabel("Ready")
        self.status_bar.addWidget(self.status_label)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setMaximumWidth(200)
        self.progress_bar.setVisible(False)
        self.status_bar.addWidget(self.progress_bar)
        
        # 创建系统托盘图标
        if QSystemTrayIcon.isSystemTrayAvailable():
            self.create_system_tray()
        
        # 创建停靠窗口
        self.create_dock_windows()
    
    def create_menubar(self):
        menubar = self.menuBar()
        
        # 文件菜单
        file_menu = menubar.addMenu("File")
        
        new_action = QAction("New Project", self)
        new_action.triggered.connect(self.new_project)
        file_menu.addAction(new_action)
        
        import_action = QAction("Import Data", self)
        import_action.triggered.connect(self.import_data)
        file_menu.addAction(import_action)
        
        export_action = QAction("Export Report", self)
        export_action.triggered.connect(self.export_report)
        file_menu.addAction(export_action)
        
        backup_action = QAction("Backup Database", self)
        backup_action.triggered.connect(self.backup_database)
        file_menu.addAction(backup_action)
        
        restore_action = QAction("Restore Database", self)
        restore_action.triggered.connect(self.restore_database)
        file_menu.addAction(restore_action)
        
        exit_action = QAction("Exit", self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # 编辑菜单
        edit_menu = menubar.addMenu("Edit")
        
        preferences_action = QAction("Preferences", self)
        preferences_action.triggered.connect(self.show_preferences)
        edit_menu.addAction(preferences_action)
        
        # 视图菜单
        view_menu = menubar.addMenu("View")
        
        refresh_action = QAction("Refresh", self)
        refresh_action.triggered.connect(self.refresh_data)
        view_menu.addAction(refresh_action)
        
        # 工具菜单
        tools_menu = menubar.addMenu("Tools")
        
        resource_action = QAction("Resource Manager", self)
        resource_action.triggered.connect(self.show_resource_manager)
        tools_menu.addAction(resource_action)
        
        risk_action = QAction("Risk Manager", self)
        risk_action.triggered.connect(self.show_risk_manager)
        tools_menu.addAction(risk_action)
        
        # 帮助菜单
        help_menu = menubar.addMenu("Help")
        
        about_action = QAction("About", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
    
    def create_toolbar(self):
        # 主工具栏
        main_toolbar = QToolBar("Main Toolbar")
        main_toolbar.setIconSize(QSize(24, 24))
        self.addToolBar(main_toolbar)
        
        new_project_action = QAction(QIcon("icons/new.png"), "New Project", self)
        new_project_action.triggered.connect(self.new_project)
        main_toolbar.addAction(new_project_action)
        
        save_action = QAction(QIcon("icons/save.png"), "Save", self)
        save_action.triggered.connect(self.save_project)
        main_toolbar.addAction(save_action)
        
        main_toolbar.addSeparator()
        
        chart_action = QAction(QIcon("icons/chart.png"), "Generate Chart", self)
        chart_action.triggered.connect(self.generate_chart)
        main_toolbar.addAction(chart_action)
        
        report_action = QAction(QIcon("icons/report.png"), "Generate Report", self)
        report_action.triggered.connect(self.generate_report)
        main_toolbar.addAction(report_action)
        
        main_toolbar.addSeparator()
        
        analytics_action = QAction(QIcon("icons/analytics.png"), "Advanced Analytics", self)
        analytics_action.triggered.connect(self.run_advanced_analytics)
        main_toolbar.addAction(analytics_action)
        
        # 分析工具栏
        analytics_toolbar = QToolBar("Analytics Toolbar")
        analytics_toolbar.setIconSize(QSize(20, 20))
        self.addToolBar(analytics_toolbar)
        
        forecast_action = QAction(QIcon("icons/forecast.png"), "Forecast", self)
        forecast_action.triggered.connect(lambda: self.run_analysis("forecast"))
        analytics_toolbar.addAction(forecast_action)
        
        cluster_action = QAction(QIcon("icons/cluster.png"), "Clustering", self)
        cluster_action.triggered.connect(lambda: self.run_analysis("clustering"))
        analytics_toolbar.addAction(cluster_action)
        
        sentiment_action = QAction(QIcon("icons/sentiment.png"), "Sentiment Analysis", self)
        sentiment_action.triggered.connect(lambda: self.run_analysis("sentiment"))
        analytics_toolbar.addAction(sentiment_action)
    
    def create_dock_windows(self):
        # 实时数据停靠窗口
        realtime_dock = QDockWidget("Real-time Monitor", self)
        realtime_dock.setAllowedAreas(Qt.LeftDockWidgetArea | Qt.RightDockWidgetArea)
        
        realtime_widget = QWidget()
        realtime_layout = QVBoxLayout()
        
        self.realtime_data_label = QLabel("No data available")
        self.realtime_data_label.setWordWrap(True)
        realtime_layout.addWidget(self.realtime_data_label)
        
        realtime_widget.setLayout(realtime_layout)
        realtime_dock.setWidget(realtime_widget)
        
        self.addDockWidget(Qt.RightDockWidgetArea, realtime_dock)
        
        # 任务列表停靠窗口
        tasks_dock = QDockWidget("Recent Tasks", self)
        tasks_dock.setAllowedAreas(Qt.LeftDockWidgetArea | Qt.RightDockWidgetArea)
        
        tasks_widget = QWidget()
        tasks_layout = QVBoxLayout()
        
        self.tasks_list = QListWidget()
        tasks_layout.addWidget(self.tasks_list)
        
        tasks_widget.setLayout(tasks_layout)
        tasks_dock.setWidget(tasks_widget)
        
        self.addDockWidget(Qt.RightDockWidgetArea, tasks_dock)
    
    def create_system_tray(self):
        self.tray_icon = QSystemTrayIcon(self)
        self.tray_icon.setIcon(self.style().standardIcon(QStyle.SP_ComputerIcon))
        
        tray_menu = QMenu()
        
        show_action = tray_menu.addAction("Show")
        show_action.triggered.connect(self.show)
        
        hide_action = tray_menu.addAction("Hide")
        hide_action.triggered.connect(self.hide)
        
        quit_action = tray_menu.addAction("Quit")
        quit_action.triggered.connect(QApplication.quit)
        
        self.tray_icon.setContextMenu(tray_menu)
        self.tray_icon.show()
        self.tray_icon.activated.connect(self.tray_icon_activated)
    
    def tray_icon_activated(self, reason):
        if reason == QSystemTrayIcon.DoubleClick:
            if self.isVisible():
                self.hide()
            else:
                self.show()
                self.activateWindow()
    
    def create_dashboard_tab(self):
        dashboard_tab = QWidget()
        layout = QVBoxLayout(dashboard_tab)
        
        # 创建仪表板部件
        summary_group = QGroupBox("Project Summary")
        summary_layout = QHBoxLayout()
        
        # 项目状态卡片
        status_cards = [
            ("Total Projects", "12", "#3498db"),
            ("Active Projects", "7", "#2ecc71"),
            ("On Hold", "3", "#f39c12"),
            ("Completed", "2", "#95a5a6"),
            ("At Risk", "2", "#e74c3c")
        ]
        
        for title, value, color in status_cards:
            card = QWidget()
            card.setStyleSheet(f"background-color: {color}; color: white; border-radius: 5px; padding: 10px;")
            card_layout = QVBoxLayout()
            card_title = QLabel(title)
            card_title.setStyleSheet("font-weight: bold; font-size: 12px;")
            card_value = QLabel(value)
            card_value.setStyleSheet("font-size: 24px; font-weight: bold;")
            card_layout.addWidget(card_title)
            card_layout.addWidget(card_value)
            card.setLayout(card_layout)
            card.setFixedSize(120, 80)
            summary_layout.addWidget(card)
        
        summary_group.setLayout(summary_layout)
        layout.addWidget(summary_group)
        
        # 图表区域
        chart_group = QGroupBox("Performance Metrics")
        chart_layout = QHBoxLayout()
        
        self.dashboard_chart = AdvancedMplCanvas(self, width=5, height=4, dpi=100)
        chart_layout.addWidget(self.dashboard_chart)
        
        # 模拟数据
        data = {
            'Completed': {'x': ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun'], 'y': [2, 3, 1, 4, 2, 3]},
            'Delayed': {'x': ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun'], 'y': [1, 0, 2, 1, 0, 1]}
        }
        self.dashboard_chart.plot_bar(data, title="Monthly Project Status", 
                                     xlabel="Month", ylabel="Count", stacked=True)
        
        chart_group.setLayout(chart_layout)
        layout.addWidget(chart_group)
        
        # 项目进度区域
        progress_group = QGroupBox("Project Progress")
        progress_layout = QVBoxLayout()
        
        self.progress_table = QTableWidget()
        self.progress_table.setColumnCount(4)
        self.progress_table.setHorizontalHeaderLabels(["Project", "Status", "Progress", "Health"])
        self.progress_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        
        # 添加示例数据
        self.progress_table.setRowCount(5)
        projects = ["Website Redesign", "Mobile App", "CRM Integration", "Data Migration", "UI/UX Update"]
        statuses = ["Active", "Active", "On Hold", "Active", "Completed"]
        progresses = [75, 45, 30, 90, 100]
        healths = ["Good", "At Risk", "Critical", "Excellent", "Completed"]
        
        for i, (project, status, progress, health) in enumerate(zip(projects, statuses, progresses, healths)):
            self.progress_table.setItem(i, 0, QTableWidgetItem(project))
            self.progress_table.setItem(i, 1, QTableWidgetItem(status))
            self.progress_table.setItem(i, 2, QTableWidgetItem(f"{progress}%"))
            
            health_item = QTableWidgetItem(health)
            if health == "Good":
                health_item.setBackground(QColor(46, 204, 113))
            elif health == "At Risk":
                health_item.setBackground(QColor(241, 196, 15))
            elif health == "Critical":
                health_item.setBackground(QColor(231, 76, 60))
            elif health == "Excellent":
                health_item.setBackground(QColor(39, 174, 96))
            elif health == "Completed":
                health_item.setBackground(QColor(149, 165, 166))
            
            health_item.setForeground(QColor(255, 255, 255))
            self.progress_table.setItem(i, 3, health_item)
        
        progress_layout.addWidget(self.progress_table)
        progress_group.setLayout(progress_layout)
        layout.addWidget(progress_group)
        
        self.tab_widget.addTab(dashboard_tab, "Dashboard")
    
    def create_projects_tab(self):
        projects_tab = QWidget()
        layout = QHBoxLayout(projects_tab)
        
        # 左侧项目列表
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        
        projects_list_label = QLabel("Projects")
        projects_list_label.setStyleSheet("font-weight: bold; font-size: 14px;")
        left_layout.addWidget(projects_list_label)
        
        # 项目过滤器
        filter_layout = QHBoxLayout()
        filter_layout.addWidget(QLabel("Filter:"))
        
        self.status_filter = QComboBox()
        self.status_filter.addItems(["All", "Planning", "Active", "On Hold", "Completed", "Cancelled"])
        self.status_filter.currentTextChanged.connect(self.filter_projects)
        filter_layout.addWidget(self.status_filter)
        
        self.priority_filter = QComboBox()
        self.priority_filter.addItems(["All", "1 - Critical", "2 - High", "3 - Medium", "4 - Low", "5 - Minimal"])
        self.priority_filter.currentTextChanged.connect(self.filter_projects)
        filter_layout.addWidget(self.priority_filter)
        
        left_layout.addLayout(filter_layout)
        
        self.projects_tree = QTreeWidget()
        self.projects_tree.setHeaderLabels(["ID", "Name", "Status", "Priority"])
        self.projects_tree.itemClicked.connect(self.on_project_selected)
        self.projects_tree.setColumnWidth(0, 50)
        self.projects_tree.setColumnWidth(1, 200)
        self.projects_tree.setColumnWidth(2, 100)
        self.projects_tree.setColumnWidth(3, 80)
        left_layout.addWidget(self.projects_tree)
        
        # 右侧项目详情
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        
        details_label = QLabel("Project Details")
        details_label.setStyleSheet("font-weight: bold; font-size: 14px;")
        right_layout.addWidget(details_label)
        
        self.project_details = AdvancedProjectDetailsWidget(db_manager=self.db_manager)
        right_layout.addWidget(self.project_details)
        
        # 按钮区域
        button_layout = QHBoxLayout()
        save_btn = QPushButton("Save")
        save_btn.clicked.connect(self.save_project)
        delete_btn = QPushButton("Delete")
        delete_btn.clicked.connect(self.delete_project)
        gantt_btn = QPushButton("View Gantt Chart")
        gantt_btn.clicked.connect(self.show_gantt_chart)
        button_layout.addWidget(save_btn)
        button_layout.addWidget(delete_btn)
        button_layout.addWidget(gantt_btn)
        right_layout.addLayout(button_layout)
        
        # 任务列表
        tasks_label = QLabel("Tasks")
        tasks_label.setStyleSheet("font-weight: bold; font-size: 14px;")
        right_layout.addWidget(tasks_label)
        
        self.tasks_table = QTableWidget()
        self.tasks_table.setColumnCount(7)
        self.tasks_table.setHorizontalHeaderLabels(["ID", "Name", "Assignee", "Status", "Priority", "Due Date", "Progress"])
        self.tasks_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        right_layout.addWidget(self.tasks_table)
        
        # 分割左右区域
        splitter = QSplitter(Qt.Horizontal)
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setSizes([300, 900])
        
        layout.addWidget(splitter)
        self.tab_widget.addTab(projects_tab, "Projects")
        
        # 加载项目数据
        self.load_projects()
    
    def create_team_tab(self):
        team_tab = QWidget()
        layout = QVBoxLayout(team_tab)
        
        # 团队成员表格
        self.team_table = QTableWidget()
        self.team_table.setColumnCount(10)
        self.team_table.setHorizontalHeaderLabels([
            "ID", "Name", "Role", "Department", "Skills", "Workload", 
            "Performance", "Hourly Rate", "Available Hours", "Contact"
        ])
        self.team_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        
        layout.addWidget(self.team_table)
        
        # 团队分析区域
        analysis_group = QGroupBox("Team Analysis")
        analysis_layout = QHBoxLayout()
        
        # 工作负载图表
        workload_chart = AdvancedMplCanvas(self, width=4, height=3)
        analysis_layout.addWidget(workload_chart)
        
        # 技能分布图表
        skills_chart = AdvancedMplCanvas(self, width=4, height=3)
        analysis_layout.addWidget(skills_chart)
        
        analysis_group.setLayout(analysis_layout)
        layout.addWidget(analysis_group)
        
        # 加载团队数据
        self.load_team_data()
        
        # 修正数据格式 - 使用字典格式
        workload_data = {
            'Team Workload': {
                'x': ['John Doe', 'Jane Smith', 'Bob Johnson', 'Alice Brown', 'Charlie Wilson'],
                'y': [75, 60, 90, 45, 80]
            }
        }
        workload_chart.plot_bar(workload_data, title="Team Workload", xlabel="Team Member", ylabel="Workload (%)")
        
        skills_data = {
            'Team Skills': {
                'x': ['Python', 'JavaScript', 'SQL', 'Project Management', 'UI/UX Design'],
                'y': [8, 6, 7, 9, 5]
            }
        }
        skills_chart.plot_bar(skills_data, title="Team Skills", xlabel="Skill", ylabel="Proficiency")
        
        self.tab_widget.addTab(team_tab, "Team")
    
    def create_analytics_tab(self):
        analytics_tab = QWidget()
        layout = QVBoxLayout(analytics_tab)
        
        # 分析控制区域
        control_group = QGroupBox("Analysis Controls")
        control_layout = QHBoxLayout()
        
        self.analysis_type_combo = QComboBox()
        self.analysis_type_combo.addItems([
            "Trend Analysis", 
            "Correlation Analysis", 
            "Forecasting",
            "Clustering",
            "Sentiment Analysis"
        ])
        control_layout.addWidget(QLabel("Analysis Type:"))
        control_layout.addWidget(self.analysis_type_combo)
        
        self.param1_label = QLabel("Periods:")
        self.param1_spin = QSpinBox()
        self.param1_spin.setRange(1, 365)
        self.param1_spin.setValue(30)
        control_layout.addWidget(self.param1_label)
        control_layout.addWidget(self.param1_spin)
        
        run_analysis_btn = QPushButton("Run Analysis")
        run_analysis_btn.clicked.connect(self.run_advanced_analytics)
        control_layout.addWidget(run_analysis_btn)
        
        control_layout.addStretch()
        
        control_group.setLayout(control_layout)
        layout.addWidget(control_group)
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)
        
        # 状态标签
        self.analysis_status_label = QLabel("")
        layout.addWidget(self.analysis_status_label)
        
        # 图表区域
        chart_group = QGroupBox("Analysis Results")
        chart_layout = QVBoxLayout()
        
        self.analytics_chart = AdvancedMplCanvas(self, width=8, height=6, dpi=100)
        chart_layout.addWidget(self.analytics_chart)
        
        chart_group.setLayout(chart_layout)
        layout.addWidget(chart_group)
        
        # 结果表格区域
        results_group = QGroupBox("Analysis Details")
        results_layout = QVBoxLayout()
        
        self.results_table = QTableWidget()
        self.results_table.setColumnCount(5)
        self.results_table.setHorizontalHeaderLabels(["Metric", "Value", "Target", "Variance", "Status"])
        results_layout.addWidget(self.results_table)
        
        results_group.setLayout(results_layout)
        layout.addWidget(results_group)
        
        self.tab_widget.addTab(analytics_tab, "Analytics")
    
    def create_reports_tab(self):
        reports_tab = QWidget()
        layout = QVBoxLayout(reports_tab)
        
        # 报告类型选择
        report_type_group = QGroupBox("Report Configuration")
        report_type_layout = QGridLayout()
        
        report_type_layout.addWidget(QLabel("Report Type:"), 0, 0)
        self.report_type_combo = QComboBox()
        self.report_type_combo.addItems([
            "Project Status Report",
            "Team Performance Report",
            "Financial Summary",
            "Risk Assessment Report",
            "Resource Utilization Report",
            "Custom Report"
        ])
        report_type_layout.addWidget(self.report_type_combo, 0, 1)
        
        report_type_layout.addWidget(QLabel("Time Period:"), 1, 0)
        self.report_period_combo = QComboBox()
        self.report_period_combo.addItems([
            "Last Week",
            "Last Month",
            "Last Quarter",
            "Last Year",
            "Custom Range"
        ])
        report_type_layout.addWidget(self.report_period_combo, 1, 1)
        
        report_type_layout.addWidget(QLabel("Format:"), 2, 0)
        self.report_format_combo = QComboBox()
        self.report_format_combo.addItems(["HTML", "PDF", "Word", "Excel"])
        report_type_layout.addWidget(self.report_format_combo, 2, 1)
        
        report_type_layout.addWidget(QLabel("Include Charts:"), 3, 0)
        self.include_charts_check = QCheckBox()
        self.include_charts_check.setChecked(True)
        report_type_layout.addWidget(self.include_charts_check, 3, 1)
        
        report_type_group.setLayout(report_type_layout)
        layout.addWidget(report_type_group)
        
        # 按钮区域
        button_layout = QHBoxLayout()
        generate_btn = QPushButton("Generate Report")
        generate_btn.clicked.connect(self.generate_report)
        button_layout.addWidget(generate_btn)
        
        export_btn = QPushButton("Export Report")
        export_btn.clicked.connect(self.export_report)
        button_layout.addWidget(export_btn)
        
        schedule_btn = QPushButton("Schedule Report")
        schedule_btn.clicked.connect(self.schedule_report)
        button_layout.addWidget(schedule_btn)
        
        layout.addLayout(button_layout)
        
        # 报告预览区域
        preview_group = QGroupBox("Report Preview")
        preview_layout = QVBoxLayout()
        
        self.report_preview = QTextEdit()
        self.report_preview.setReadOnly(True)
        preview_layout.addWidget(self.report_preview)
        
        preview_group.setLayout(preview_layout)
        layout.addWidget(preview_group)
        
        self.tab_widget.addTab(reports_tab, "Reports")
    
    def create_risks_tab(self):
        risks_tab = QWidget()
        layout = QVBoxLayout(risks_tab)
        
        # 风险概述
        overview_group = QGroupBox("Risk Overview")
        overview_layout = QHBoxLayout()
        
        risk_cards = [
            ("Total Risks", "15", "#3498db"),
            ("High Priority", "5", "#e74c3c"),
            ("Medium Priority", "7", "#f39c12"),
            ("Low Priority", "3", "#2ecc71")
        ]
        
        for title, value, color in risk_cards:
            card = QWidget()
            card.setStyleSheet(f"background-color: {color}; color: white; border-radius: 5px; padding: 10px;")
            card_layout = QVBoxLayout()
            card_title = QLabel(title)
            card_title.setStyleSheet("font-weight: bold; font-size: 12px;")
            card_value = QLabel(value)
            card_value.setStyleSheet("font-size: 24px; font-weight: bold;")
            card_layout.addWidget(card_title)
            card_layout.addWidget(card_value)
            card.setLayout(card_layout)
            card.setFixedSize(120, 80)
            overview_layout.addWidget(card)
        
        overview_group.setLayout(overview_layout)
        layout.addWidget(overview_group)
        
        # 风险表格
        risks_group = QGroupBox("Risk Register")
        risks_layout = QVBoxLayout()
        
        self.risks_table = QTableWidget()
        self.risks_table.setColumnCount(9)
        self.risks_table.setHorizontalHeaderLabels([
            "ID", "Project", "Description", "Impact", "Probability", "Mitigation Plan", "Status", "Assigned To", "Due Date"
        ])
        self.risks_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        risks_layout.addWidget(self.risks_table)
        
        risks_group.setLayout(risks_layout)
        layout.addWidget(risks_group)
        
        # 加载风险数据
        self.load_risks_data()
        
        self.tab_widget.addTab(risks_tab, "Risks")
    
    def create_resources_tab(self):
        resources_tab = QWidget()
        layout = QVBoxLayout(resources_tab)
        
        # 资源概述
        overview_group = QGroupBox("Resource Overview")
        overview_layout = QHBoxLayout()
        
        resource_cards = [
            ("Total Resources", "28", "#3498db"),
            ("Allocated", "19", "#2ecc71"),
            ("Available", "9", "#f39c12"),
            ("Overallocated", "3", "#e74c3c")
        ]
        
        for title, value, color in resource_cards:
            card = QWidget()
            card.setStyleSheet(f"background-color: {color}; color: white; border-radius: 5px; padding: 10px;")
            card_layout = QVBoxLayout()
            card_title = QLabel(title)
            card_title.setStyleSheet("font-weight: bold; font-size: 12px;")
            card_value = QLabel(value)
            card_value.setStyleSheet("font-size: 24px; font-weight: bold;")
            card_layout.addWidget(card_title)
            card_layout.addWidget(card_value)
            card.setLayout(card_layout)
            card.setFixedSize(120, 80)
            overview_layout.addWidget(card)
        
        overview_group.setLayout(overview_layout)
        layout.addWidget(overview_group)
        
        # 资源表格
        resources_group = QGroupBox("Resource Allocation")
        resources_layout = QVBoxLayout()
        
        self.resources_table = QTableWidget()
        self.resources_table.setColumnCount(7)
        self.resources_table.setHorizontalHeaderLabels([
            "ID", "Name", "Type", "Cost", "Availability", "Assigned To", "Project"
        ])
        self.resources_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        resources_layout.addWidget(self.resources_table)
        
        resources_group.setLayout(resources_layout)
        layout.addWidget(resources_group)
        
        # 加载资源数据
        self.load_resources_data()
        
        self.tab_widget.addTab(resources_tab, "Resources")
    
    def create_documents_tab(self):
        documents_tab = QWidget()
        layout = QVBoxLayout(documents_tab)
        
        # 文档管理区域
        docs_group = QGroupBox("Document Management")
        docs_layout = QVBoxLayout()
        
        # 工具栏
        docs_toolbar = QToolBar()
        upload_action = QAction("Upload Document", self)
        download_action = QAction("Download Document", self)
        delete_action = QAction("Delete Document", self)
        
        docs_toolbar.addAction(upload_action)
        docs_toolbar.addAction(download_action)
        docs_toolbar.addAction(delete_action)
        
        docs_layout.addWidget(docs_toolbar)
        
        # 文档表格
        self.documents_table = QTableWidget()
        self.documents_table.setColumnCount(7)
        self.documents_table.setHorizontalHeaderLabels([
            "ID", "Project", "Name", "Type", "Version", "Upload Date", "Uploaded By"
        ])
        self.documents_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        docs_layout.addWidget(self.documents_table)
        
        docs_group.setLayout(docs_layout)
        layout.addWidget(docs_group)
        
        # 加载文档数据
        self.load_documents_data()
        
        self.tab_widget.addTab(documents_tab, "Documents")
    
    def load_projects(self, status_filter=None, priority_filter=None):
        self.projects_tree.clear()
        
        if status_filter and status_filter != "All":
            projects = self.db_manager.get_projects(status_filter)
        else:
            projects = self.db_manager.get_projects()
        
        for project in projects:
            item = QTreeWidgetItem([str(project[0]), project[1], project[3], str(project[4])])
            
            # 根据状态设置颜色
            if project[3] == "Completed":
                item.setForeground(2, QColor(46, 204, 113))
            elif project[3] == "Active":
                item.setForeground(2, QColor(52, 152, 219))
            elif project[3] == "On Hold":
                item.setForeground(2, QColor(243, 156, 18))
            elif project[3] == "Cancelled":
                item.setForeground(2, QColor(231, 76, 60))
            
            # 根据优先级设置颜色
            if project[4] == 1:
                item.setForeground(3, QColor(231, 76, 60))
            elif project[4] == 2:
                item.setForeground(3, QColor(230, 126, 34))
            elif project[4] == 3:
                item.setForeground(3, QColor(241, 196, 15))
            
            self.projects_tree.addTopLevelItem(item)
    
    def filter_projects(self):
        status_filter = self.status_filter.currentText()
        priority_filter = self.priority_filter.currentText()
        
        # 简化实现，实际应用中需要更复杂的过滤逻辑
        self.load_projects(status_filter if status_filter != "All" else None)
    
    def load_team_data(self):
        team_members = self.db_manager.get_team_members()
        self.team_table.setRowCount(len(team_members))
        
        for row, member in enumerate(team_members):
            for col, value in enumerate(member):
                self.team_table.setItem(row, col, QTableWidgetItem(str(value)))
        
        self.team_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
    
    def load_risks_data(self):
        risks = self.db_manager.get_risks()
        self.risks_table.setRowCount(len(risks))
        
        for row, risk in enumerate(risks):
            for col, value in enumerate(risk):
                self.risks_table.setItem(row, col, QTableWidgetItem(str(value)))
        
        self.risks_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
    
    def load_resources_data(self):
        resources = self.db_manager.get_resources()
        self.resources_table.setRowCount(len(resources))
        
        for row, resource in enumerate(resources):
            for col, value in enumerate(resource):
                self.resources_table.setItem(row, col, QTableWidgetItem(str(value)))
        
        self.resources_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
    
    def load_documents_data(self):
        documents = self.db_manager.get_documents()
        self.documents_table.setRowCount(len(documents))
        
        for row, document in enumerate(documents):
            for col, value in enumerate(document):
                self.documents_table.setItem(row, col, QTableWidgetItem(str(value)))
        
        self.documents_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
    
    def on_project_selected(self, item, column):
        project_id = int(item.text(0))
        self.current_project_id = project_id
        
        # 获取项目详情
        projects = self.db_manager.get_projects()
        for project in projects:
            if project[0] == project_id:
                project_data = {
                    'name': project[1],
                    'description': project[2],
                    'status': project[3],
                    'priority': project[4],
                    'start_date': project[5],
                    'end_date': project[6],
                    'budget': project[7],
                    'actual_cost': project[8],
                    'progress': project[9],
                    'risk_level': project[13],
                    'tags': project[12],
                    'stakeholders': project[14]
                }
                self.project_details = AdvancedProjectDetailsWidget(project_data, self.db_manager)
                break
        
        # 加载项目任务
        tasks = self.db_manager.get_tasks(project_id)
        self.tasks_table.setRowCount(len(tasks))
        
        for row, task in enumerate(tasks):
            for col, value in enumerate(task):
                self.tasks_table.setItem(row, col, QTableWidgetItem(str(value)))
        
        self.tasks_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
    
    def new_project(self):
        self.project_details = AdvancedProjectDetailsWidget(db_manager=self.db_manager)
        self.current_project_id = None
        QMessageBox.information(self, "New Project", "Ready to create a new project")
    
    def save_project(self):
        project_data = self.project_details.get_data()
        
        if self.current_project_id:
            # 更新现有项目
            query = """
                UPDATE projects 
                SET name=?, description=?, status=?, priority=?, start_date=?, end_date=?, 
                budget=?, actual_cost=?, progress=?, risk_level=?, tags=?, stakeholders=?
                WHERE id=?
            """
            params = (
                project_data['name'], project_data['description'], project_data['status'],
                project_data['priority'], project_data['start_date'], project_data['end_date'],
                project_data['budget'], project_data['actual_cost'], project_data['progress'],
                project_data['risk_level'], project_data['tags'], project_data['stakeholders'],
                self.current_project_id
            )
            result = self.db_manager.execute_query(query, params, fetch=False)
            if isinstance(result, str) and result.startswith("Error"):
                QMessageBox.critical(self, "Error", f"Failed to update project: {result}")
            else:
                QMessageBox.information(self, "Success", "Project updated successfully")
        else:
            # 插入新项目
            query = """
                INSERT INTO projects (name, description, status, priority, start_date, end_date, 
                budget, actual_cost, progress, created_date, last_modified, risk_level, tags, stakeholders)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, datetime('now'), datetime('now'), ?, ?, ?)
            """
            params = (
                project_data['name'], project_data['description'], project_data['status'],
                project_data['priority'], project_data['start_date'], project_data['end_date'],
                project_data['budget'], project_data['actual_cost'], project_data['progress'],
                project_data['risk_level'], project_data['tags'], project_data['stakeholders']
            )
            result = self.db_manager.execute_query(query, params, fetch=False)
            if isinstance(result, str) and result.startswith("Error"):
                QMessageBox.critical(self, "Error", f"Failed to create project: {result}")
            else:
                QMessageBox.information(self, "Success", "Project created successfully")
        
        self.load_projects()
    
    def delete_project(self):
        if self.current_project_id:
            reply = QMessageBox.question(self, "Confirm Delete", 
                                        "Are you sure you want to delete this project?",
                                        QMessageBox.Yes | QMessageBox.No)
            
            if reply == QMessageBox.Yes:
                result = self.db_manager.execute_query("DELETE FROM projects WHERE id=?", (self.current_project_id,), fetch=False)
                if isinstance(result, str) and result.startswith("Error"):
                    QMessageBox.critical(self, "Error", f"Failed to delete project: {result}")
                else:
                    self.db_manager.execute_query("DELETE FROM tasks WHERE project_id=?", (self.current_project_id,), fetch=False)
                    QMessageBox.information(self, "Success", "Project deleted successfully")
                    self.load_projects()
                    self.current_project_id = None
                    self.project_details = AdvancedProjectDetailsWidget(db_manager=self.db_manager)
        else:
            QMessageBox.warning(self, "Warning", "No project selected")
    
    def run_advanced_analytics(self):
        analysis_type = self.analysis_type_combo.currentText().lower().replace(" ", "_")
        
        # 模拟数据
        if analysis_type == "trend_analysis":
            data = pd.DataFrame({
                'date': pd.date_range('2023-01-01', periods=100, freq='D'),
                'value': np.random.randn(100).cumsum() + 50
            })
            self.analysis_thread = AdvancedAnalysisThread(data, "trend")
        elif analysis_type == "forecasting":
            data = pd.DataFrame({
                'date': pd.date_range('2023-01-01', periods=100, freq='D'),
                'value': np.random.randn(100).cumsum() + 50
            })
            self.analysis_thread = AdvancedAnalysisThread(data, "forecast", {'periods': self.param1_spin.value()})
        elif analysis_type == "clustering":
            data = pd.DataFrame({
                'feature1': np.random.randn(100) * 10 + 50,
                'feature2': np.random.randn(100) * 5 + 30
            })
            self.analysis_thread = AdvancedAnalysisThread(data, "clustering", {'n_clusters': 3})
        elif analysis_type == "sentiment_analysis":
            data = pd.DataFrame({
                'text': [
                    "This project is going great! The team is doing an excellent job.",
                    "I'm concerned about the timeline. We might need more resources.",
                    "The client feedback has been very positive so far.",
                    "We're experiencing some technical difficulties that need to be addressed.",
                    "The budget is on track and the deliverables are meeting expectations."
                ]
            })
            self.analysis_thread = AdvancedAnalysisThread(data, "sentiment")
        
        self.progress_bar.setVisible(True)
        self.analysis_thread.progress_signal.connect(self.update_analysis_progress)
        self.analysis_thread.result_signal.connect(self.on_analysis_complete)
        self.analysis_thread.start()
    
    def update_analysis_progress(self, value, message):
        self.progress_bar.setValue(value)
        self.analysis_status_label.setText(message)
    
    def on_analysis_complete(self, result, analysis_type):
        self.progress_bar.setVisible(False)
        self.progress_bar.setValue(0)
        
        if isinstance(result, str) and result.startswith("Error"):
            QMessageBox.critical(self, "Analysis Error", result)
            return
        
        if analysis_type == "trend":
            self.analytics_chart.plot_line(
                {'x': result['date'], 'y': result['value']},
                title="Trend Analysis",
                xlabel="Date",
                ylabel="Value"
            )
        elif analysis_type == "forecast":
            # 绘制历史数据和预测
            historical_dates = pd.date_range('2023-01-01', periods=100, freq='D')
            forecast_dates = result['dates']
            
            self.analytics_chart.fig.clear()
            ax = self.analytics_chart.fig.add_subplot(111)
            
            # 绘制历史数据
            ax.plot(historical_dates, self.analysis_thread.data['value'], label='Historical')
            
            # 绘制预测数据
            ax.plot(forecast_dates, result['forecast'], label='Forecast', linestyle='--')
            
            ax.set_title("Forecast Analysis")
            ax.set_xlabel("Date")
            ax.set_ylabel("Value")
            ax.legend()
            
            self.analytics_chart.fig.tight_layout()
            self.analytics_chart.draw()
            
            # 更新结果表格
            self.results_table.setRowCount(2)
            self.results_table.setItem(0, 0, QTableWidgetItem("Model Score"))
            self.results_table.setItem(0, 1, QTableWidgetItem(f"{result['model_score']:.4f}"))
            self.results_table.setItem(1, 0, QTableWidgetItem("Forecast Periods"))
            self.results_table.setItem(1, 1, QTableWidgetItem(str(len(result['forecast']))))
        
        elif analysis_type == "clustering":
            # 绘制聚类结果
            self.analytics_chart.fig.clear()
            ax = self.analytics_chart.fig.add_subplot(111)
            
            scatter = ax.scatter(
                self.analysis_thread.data['feature1'],
                self.analysis_thread.data['feature2'],
                c=result['clusters'],
                cmap='viridis'
            )
            
            ax.set_title("Clustering Analysis")
            ax.set_xlabel("Feature 1")
            ax.set_ylabel("Feature 2")
            self.analytics_chart.fig.colorbar(scatter)
            
            self.analytics_chart.fig.tight_layout()
            self.analytics_chart.draw()
            
            # 更新结果表格
            self.results_table.setRowCount(3)
            self.results_table.setItem(0, 0, QTableWidgetItem("Number of Clusters"))
            self.results_table.setItem(0, 1, QTableWidgetItem(str(len(np.unique(result['clusters'])))))
            self.results_table.setItem(1, 0, QTableWidgetItem("Inertia"))
            self.results_table.setItem(1, 1, QTableWidgetItem(f"{result['inertia']:.4f}"))
        
        elif analysis_type == "sentiment":
            # 显示情感分析结果
            sentiment_counts = result['sentiment'].value_counts()
            
            self.analytics_chart.plot_pie(
                {'x': sentiment_counts.index.tolist(), 'y': sentiment_counts.values.tolist()},
                title="Sentiment Analysis"
            )
            
            # 更新结果表格
            self.results_table.setRowCount(len(sentiment_counts))
            for i, (sentiment, count) in enumerate(sentiment_counts.items()):
                self.results_table.setItem(i, 0, QTableWidgetItem(sentiment))
                self.results_table.setItem(i, 1, QTableWidgetItem(str(count)))
        
        self.analysis_status_label.setText("Analysis completed successfully!")
    
    def generate_report(self):
        report_type = self.report_type_combo.currentText()
        report_period = self.report_period_combo.currentText()
        include_charts = self.include_charts_check.isChecked()
        
        if report_type == "Project Status Report":
            report_content = self.generate_project_status_report()
        elif report_type == "Team Performance Report":
            report_content = self.generate_team_performance_report()
        elif report_type == "Risk Assessment Report":
            report_content = self.generate_risk_assessment_report()
        else:
            report_content = "Report generation not implemented for this type yet."
        
        self.report_preview.setHtml(report_content)
    
    def generate_project_status_report(self):
        projects = self.db_manager.get_projects()
        
        html = """
        <html>
        <head>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                h1 { color: #2c3e50; }
                h2 { color: #34495e; }
                table { border-collapse: collapse; width: 100%; margin-bottom: 20px; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #f2f2f2; }
                .completed { background-color: #d4edda; }
                .active { background-color: #d1ecf1; }
                .on-hold { background-color: #fff3cd; }
                .cancelled { background-color: #f8d7da; }
                .critical { color: #dc3545; font-weight: bold; }
                .high { color: #fd7e14; font-weight: bold; }
                .medium { color: #ffc107; font-weight: bold; }
                .low { color: #28a745; font-weight: bold; }
            </style>
        </head>
        <body>
        """
        
        html += f"<h1>Project Status Report</h1>"
        html += f"<p>Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>"
        html += f"<p>Report Period: {self.report_period_combo.currentText()}</p>"
        
        # 项目汇总
        html += "<h2>Project Summary</h2>"
        html += "<table>"
        html += "<tr><th>Status</th><th>Count</th></tr>"
        
        status_counts = {}
        for project in projects:
            status = project[3]
            status_counts[status] = status_counts.get(status, 0) + 1
        
        for status, count in status_counts.items():
            html += f"<tr><td>{status}</td><td>{count}</td></tr>"
        
        html += "</table>"
        
        # 项目详情
        html += "<h2>Project Details</h2>"
        html += "<table>"
        html += "<tr><th>ID</th><th>Name</th><th>Status</th><th>Priority</th><th>Progress</th><th>Budget</th><th>Actual Cost</th></tr>"
        
        for project in projects:
            status_class = ""
            if project[3] == "Completed":
                status_class = "completed"
            elif project[3] == "Active":
                status_class = "active"
            elif project[3] == "On Hold":
                status_class = "on-hold"
            elif project[3] == "Cancelled":
                status_class = "cancelled"
            
            priority_class = ""
            if project[4] == 1:
                priority_class = "critical"
            elif project[4] == 2:
                priority_class = "high"
            elif project[4] == 3:
                priority_class = "medium"
            elif project[4] >= 4:
                priority_class = "low"
            
            html += f"<tr class='{status_class}'>"
            html += f"<td>{project[0]}</td>"
            html += f"<td>{project[1]}</td>"
            html += f"<td>{project[3]}</td>"
            html += f"<td class='{priority_class}'>{project[4]}</td>"
            html += f"<td>{project[9]}%</td>"
            html += f"<td>${project[7]:,.2f}</td>"
            html += f"<td>${project[8]:,.2f}</td>"
            html += "</tr>"
        
        html += "</table>"
        html += "</body></html>"
        
        return html
    
    def generate_team_performance_report(self):
        team_members = self.db_manager.get_team_members()
        
        html = """
        <html>
        <head>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                h1 { color: #2c3e50; }
                h2 { color: #34495e; }
                table { border-collapse: collapse; width: 100%; margin-bottom: 20px; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #f2f2f2; }
                .excellent { background-color: #d4edda; }
                .good { background-color: #d1ecf1; }
                .average { background-color: #fff3cd; }
                .poor { background-color: #f8d7da; }
            </style>
        </head>
        <body>
        """
        
        html += f"<h1>Team Performance Report</h1>"
        html += f"<p>Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>"
        
        # 团队汇总
        html += "<h2>Team Summary</h2>"
        html += "<table>"
        html += "<tr><th>Department</th><th>Count</th><th>Avg Performance</th></tr>"
        
        dept_stats = {}
        for member in team_members:
            dept = member[3]
            if dept not in dept_stats:
                dept_stats[dept] = {'count': 0, 'total_perf': 0}
            
            dept_stats[dept]['count'] += 1
            dept_stats[dept]['total_perf'] += member[6] or 0
        
        for dept, stats in dept_stats.items():
            avg_perf = stats['total_perf'] / stats['count'] if stats['count'] > 0 else 0
            html += f"<tr><td>{dept}</td><td>{stats['count']}</td><td>{avg_perf:.2f}</td></tr>"
        
        html += "</table>"
        
        # 成员详情
        html += "<h2>Team Member Details</h2>"
        html += "<table>"
        html += "<tr><th>ID</th><th>Name</th><th>Role</th><th>Department</th><th>Performance</th><th>Workload</th></tr>"
        
        for member in team_members:
            perf_class = ""
            perf = member[6] or 0
            if perf >= 90:
                perf_class = "excellent"
            elif perf >= 75:
                perf_class = "good"
            elif perf >= 60:
                perf_class = "average"
            else:
                perf_class = "poor"
            
            html += f"<tr class='{perf_class}'>"
            html += f"<td>{member[0]}</td>"
            html += f"<td>{member[1]}</td>"
            html += f"<td>{member[2]}</td>"
            html += f"<td>{member[3]}</td>"
            html += f"<td>{perf}</td>"
            html += f"<td>{member[5]}%</td>"
            html += "</tr>"
        
        html += "</table>"
        html += "</body></html>"
        
        return html
    
    def generate_risk_assessment_report(self):
        risks = self.db_manager.get_risks()
        
        html = """
        <html>
        <head>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                h1 { color: #2c3e50; }
                h2 { color: #34495e; }
                table { border-collapse: collapse; width: 100%; margin-bottom: 20px; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #f2f2f2; }
                .high { background-color: #f8d7da; }
                .medium { background-color: #fff3cd; }
                .low { background-color: #d4edda; }
            </style>
        </head>
        <body>
        """
        
        html += f"<h1>Risk Assessment Report</h1>"
        html += f"<p>Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>"
        
        # 风险汇总
        html += "<h2>Risk Summary</h2>"
        html += "<table>"
        html += "<tr><th>Impact Level</th><th>Count</th></tr>"
        
        impact_counts = {}
        for risk in risks:
            impact = risk[3]
            impact_counts[impact] = impact_counts.get(impact, 0) + 1
        
        for impact, count in impact_counts.items():
            html += f"<tr><td>{impact}</td><td>{count}</td></tr>"
        
        html += "</table>"
        
        # 风险详情
        html += "<h2>Risk Details</h2>"
        html += "<table>"
        html += "<tr><th>ID</th><th>Project</th><th>Description</th><th>Impact</th><th>Probability</th><th>Status</th><th>Due Date</th></tr>"
        
        for risk in risks:
            impact_class = ""
            if risk[3] == "High":
                impact_class = "high"
            elif risk[3] == "Medium":
                impact_class = "medium"
            else:
                impact_class = "low"
            
            html += f"<tr class='{impact_class}'>"
            html += f"<td>{risk[0]}</td>"
            # 获取项目名称
            project = self.db_manager.get_project(risk[1])
            project_name = project[0][1] if project else "Unknown"
            html += f"<td>{project_name}</td>"
            html += f"<td>{risk[2]}</td>"
            html += f"<td>{risk[3]}</td>"
            html += f"<td>{risk[4]}</td>"
            html += f"<td>{risk[6]}</td>"
            html += f"<td>{risk[8]}</td>"
            html += "</tr>"
        
        html += "</table>"
        html += "</body></html>"
        
        return html
    
    def export_report(self):
        file_format = self.report_format_combo.currentText().lower()
        file_path, _ = QFileDialog.getSaveFileName(
            self, 
            "Export Report", 
            "", 
            f"{file_format.upper()} Files (*.{file_format})"
        )
        
        if file_path:
            try:
                if file_format == "docx":
                    self.export_to_word(file_path)
                elif file_format == "pdf":
                    self.export_to_pdf(file_path)
                elif file_format == "html":
                    self.export_to_html(file_path)
                elif file_format == "xlsx":
                    self.export_to_excel(file_path)
                
                QMessageBox.information(self, "Success", f"Report exported to {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to export report: {str(e)}")
    
    def export_to_word(self, file_path):
        doc = Document()
        
        # 添加标题
        title = doc.add_heading('Product Manager Director System Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加生成日期
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        # 添加报告内容
        html_content = self.report_preview.toHtml()
        # 这里需要将HTML转换为Word格式（简化处理）
        # 实际应用中可能需要使用更复杂的转换库
        
        doc.add_paragraph("This is a placeholder for the actual report content.")
        
        doc.save(file_path)
    
    def export_to_pdf(self, file_path):
        # 简化实现，实际应用中可以使用wkhtmltopdf或其他库
        html_content = self.report_preview.toHtml()
        # 将HTML转换为PDF的逻辑
        pass
    
    def export_to_html(self, file_path):
        html_content = self.report_preview.toHtml()
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
    
    def export_to_excel(self, file_path):
        # 简化实现，实际应用中可以使用pandas
        pass
    
    def schedule_report(self):
        # 显示报告调度对话框
        QMessageBox.information(self, "Schedule Report", "Report scheduling feature would open here")
    
    def import_data(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, 
            "Import Data", 
            "", 
            "CSV Files (*.csv);;Excel Files (*.xlsx *.xls);;JSON Files (*.json)"
        )
        
        if file_path:
            try:
                if file_path.endswith('.csv'):
                    data = pd.read_csv(file_path)
                elif file_path.endswith(('.xlsx', '.xls')):
                    data = pd.read_excel(file_path)
                elif file_path.endswith('.json'):
                    data = pd.read_json(file_path)
                
                # 这里可以添加数据处理的逻辑
                QMessageBox.information(self, "Success", f"Data imported successfully: {len(data)} rows")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to import data: {str(e)}")
    
    def backup_database(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Backup Database",
            "",
            "Database Files (*.db *.sqlite)"
        )
        
        if file_path:
            success = self.db_manager.backup_database(file_path)
            if success:
                QMessageBox.information(self, "Success", f"Database backed up to {file_path}")
            else:
                QMessageBox.critical(self, "Error", "Failed to backup database")
    
    def restore_database(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Restore Database",
            "",
            "Database Files (*.db *.sqlite)"
        )
        
        if file_path:
            reply = QMessageBox.question(
                self,
                "Confirm Restore",
                "Are you sure you want to restore the database? This will overwrite all current data.",
                QMessageBox.Yes | QMessageBox.No
            )
            
            if reply == QMessageBox.Yes:
                success = self.db_manager.restore_database(file_path)
                if success:
                    QMessageBox.information(self, "Success", "Database restored successfully")
                    self.refresh_data()
                else:
                    QMessageBox.critical(self, "Error", "Failed to restore database")
    
    def show_resource_manager(self):
        dialog = ResourceManagerDialog(self.db_manager, self)
        dialog.exec_()
    
    def show_risk_manager(self):
        dialog = RiskManagerDialog(self.db_manager, self.current_project_id, self)
        dialog.exec_()
    
    def show_gantt_chart(self):
        if not self.current_project_id:
            QMessageBox.warning(self, "Warning", "Please select a project first")
            return
        
        # 获取项目任务
        tasks = self.db_manager.get_tasks(self.current_project_id)
        
        # 准备甘特图数据
        gantt_data = []
        for task in tasks:
            gantt_data.append({
                'name': task[2],
                'start_date': task[9],
                'end_date': task[10],
                'status': task[5]
            })
        
        # 创建甘特图对话框
        dialog = QDialog(self)
        dialog.setWindowTitle("Gantt Chart")
        dialog.setMinimumSize(1000, 600)
        
        layout = QVBoxLayout()
        gantt_widget = GanttWidget()
        gantt_widget.set_tasks(gantt_data)
        layout.addWidget(gantt_widget)
        
        dialog.setLayout(layout)
        dialog.exec_()
    
    def show_preferences(self):
        # 显示首选项对话框
        QMessageBox.information(self, "Preferences", "Preferences dialog would open here")
    
    def refresh_data(self):
        self.load_projects()
        self.load_team_data()
        self.load_risks_data()
        self.load_resources_data()
        self.load_documents_data()
        self.status_bar.showMessage("Data refreshed", 3000)
    
    def update_real_time_data(self, data):
        text = f"""
        <b>Real-time Monitoring</b><br>
        Timestamp: {data['timestamp']}<br>
        Active Users: {data['active_users']}<br>
        Server Load: {data['server_load']:.2%}<br>
        Response Time: {data['response_time']:.2f} ms<br>
        Errors: {data['errors']}
        """
        self.realtime_data_label.setText(text)
    
    def generate_chart(self):
        # 打开图表生成对话框
        QMessageBox.information(self, "Chart Generation", "Chart generation feature would open here")
    
    def show_about(self):
        about_text = """
        <h2>Advanced Product Manager Director System</h2>
        <p>Version 2.0</p>
        <p>A comprehensive tool for product managers and directors to manage projects, teams, and analytics.</p>
        <p><b>Features:</b></p>
        <ul>
            <li>Advanced Project Management</li>
            <li>Team Performance Tracking</li>
            <li>Risk Management</li>
            <li>Resource Allocation</li>
            <li>Advanced Analytics and Forecasting</li>
            <li>Real-time Monitoring</li>
            <li>Custom Report Generation</li>
        </ul>
        <p>© 2025 Product Management Solutions Inc.</p>
        """
        QMessageBox.about(self, "About", about_text)
    
    def load_settings(self):
        # 加载应用程序设置
        geometry = self.settings.value("geometry")
        if geometry:
            self.restoreGeometry(geometry)
        
        state = self.settings.value("windowState")
        if state:
            self.restoreState(state)
    
    def closeEvent(self, event):
        # 保存应用程序设置
        self.settings.setValue("geometry", self.saveGeometry())
        self.settings.setValue("windowState", self.saveState())
        
        # 停止实时监视器
        if hasattr(self, 'monitor'):
            self.monitor.stop()
            self.monitor.wait()
        
        event.accept()


# 应用程序入口
if __name__ == "__main__":
    app = QApplication(sys.argv)
    
    # 设置应用程序样式
    app.setStyle("Fusion")
    
    # 创建调色板
    palette = QPalette()
    palette.setColor(QPalette.Window, QColor(240, 240, 240))
    palette.setColor(QPalette.WindowText, Qt.black)
    app.setPalette(palette)
    
    # 设置应用程序字体
    font = QFont("Segoe UI", 10)
    app.setFont(font)
    
    window = AdvancedProductManagerDirectorSystem()
    window.show()
    
    sys.exit(app.exec_())